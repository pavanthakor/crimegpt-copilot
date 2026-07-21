"""Generate the 4 CrimeGPT .docx templates (docxtpl / Jinja placeholders).

Binary .docx are committed, but this script is the reproducible source of truth.
Run from the repo root (or anywhere):  python templates/_build_templates.py

Placeholders use docxtpl syntax:
  {{ field }}                      scalar merge field
  {%tr for x in list %} … {%tr endfor %}   table-row loop (3-row pattern:
        a for-row and endfor-row that docxtpl deletes, wrapping the content row)
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = Path(__file__).resolve().parent

# Complex-script font for Gujarati/Devanagari text. Pinned so the rendered .docx never
# falls back to the viewer's default Indic font (which prints box glyphs where Noto is not
# the OS fallback). English/Latin text keeps its ORIGINAL face: only runs that actually
# contain Indic characters get Noto in the Latin (ascii/hAnsi) slots — Latin-only runs get
# Noto in the complex-script (w:cs) slot alone. The .ttf is bundled in fonts/ and must be
# installed on the demo machine — see README "Fonts".
GUJARATI_FONT = "Noto Sans Gujarati"
# OOXML w:rFonts holds ONE face per slot, so a CSS-style fallback list cannot live inside
# the document. The primary is pinned below; if it is not installed Word substitutes the
# next available face in this documented chain (all cover the Gujarati block U+0A80–U+0AFF).
FONT_FALLBACK_CHAIN = [GUJARATI_FONT, "Nirmala UI", "Shruti", "Arial Unicode MS"]

# Placeholders whose merged value is translated to Gujarati/Hindi in a non-English doc
# (mirror of services/documents.py `_TRANSLATABLE_BY_DOC` — keep the two in sync). Their
# runs hold a Latin `{{ ... }}` tag at build time but render to Indic text, so they must be
# pinned as Indic even though _has_indic() sees only Latin now.
INDIC_BOUND_FIELDS = (
    "proceedings_narrative",
    "investigation_done",
    "pending_investigation",
    "grounds_for_custody",
    "examination_purpose",
    "complaint_narrative",
)


def _has_indic(text):
    """True if text has any Gujarati (U+0A80–U+0AFF) or Devanagari (U+0900–U+097F) char."""
    return any("ऀ" <= c <= "ॿ" or "઀" <= c <= "૿" for c in text)


def _pin_run(run):
    """Pin the Indic face on a run without disturbing Latin text.

    Indic run  -> all four slots (ascii/hAnsi/cs/eastAsia); belt-and-braces, since some
                  viewers route Indic text through the ascii slot. A run counts as Indic if
                  it already holds Indic characters OR holds a Gujarati-bound placeholder
                  that will render to Indic text.
    Latin run  -> only w:cs, so ascii/hAnsi keep the document's original Latin face.
    """
    text = run.text
    is_indic = _has_indic(text) or any(field in text for field in INDIC_BOUND_FIELDS)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    if is_indic:
        for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(slot), GUJARATI_FONT)
    else:
        rfonts.set(qn("w:cs"), GUJARATI_FONT)


def _iter_runs(doc):
    """Yield every run in body paragraphs and (single-level) table cells."""
    for para in doc.paragraphs:
        yield from para.runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield from para.runs


def _apply_fonts(doc):
    """Set the complex-script default to Noto and pin every run. Call once before saving.

    Only the complex-script (w:cs) default is changed, so even runs that docxtpl inserts
    at render time resolve Gujarati to Noto. The Latin (ascii/hAnsi) default is left
    untouched, so English keeps the document's original Latin face.
    """
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)  # w:rFonts must be the first child of w:rPr
    # Complex-script default -> Noto (drop the theme ref so the explicit face wins).
    # ascii/hAnsi/eastAsia are left as-is, preserving the original Latin default.
    if rfonts.get(qn("w:cstheme")) is not None:
        del rfonts.attrib[qn("w:cstheme")]
    rfonts.set(qn("w:cs"), GUJARATI_FONT)

    for run in _iter_runs(doc):
        _pin_run(run)


def _title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def _line(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def _looped_table(doc, headers, for_expr, content_cells, endfor="{%tr endfor %}"):
    """Build a table whose single content row repeats via docxtpl {%tr%}.

    Rows: [header] [for-tag] [content] [endfor-tag]. docxtpl removes the for/endfor
    rows and repeats the content row for each item.
    """
    n = len(headers)
    table = doc.add_table(rows=4, cols=n)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    # for-row: tag in first cell, rest blank (docxtpl drops the whole row)
    table.rows[1].cells[0].text = for_expr
    # content-row: the repeated cells
    for i, c in enumerate(content_cells):
        table.rows[2].cells[i].text = c
    # endfor-row
    table.rows[3].cells[0].text = endfor
    return table


def _sig_block(doc, left_label, right_label):
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = left_label
    t.rows[0].cells[1].text = right_label
    t.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _blank_box(doc, lines=3):
    """A single bordered cell with blank height — e.g. for a physical seal impression."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    for _ in range(lines):
        cell.add_paragraph()
    return t


# --------------------------------------------------------------------------
def build_seizure_receipt():
    """CCTNS Form IF4 — Property Seizure Memo layout (uses existing pool fields only)."""
    doc = Document()
    _title(doc, "FORM IF4 — PROPERTY SEIZURE MEMO / જપ્તી પંચનામું")
    _line(doc, "(Crime & Criminal Tracking Network and Systems · Bharatiya Nagarik Suraksha Sanhita, 2023)")
    doc.add_paragraph()

    # District / PS / Year / FIR header
    hdr = doc.add_table(rows=2, cols=2)
    hdr.style = "Table Grid"
    hdr.rows[0].cells[0].text = "District: {{ district }}"
    hdr.rows[0].cells[1].text = "Police Station: {{ police_station }}"
    hdr.rows[1].cells[0].text = "Year: {{ fir_year }}    FIR No.: {{ fir_number }}"
    hdr.rows[1].cells[1].text = "Crime / Case No.: {{ case_number }}    FIR Date: {{ fir_date }}"
    doc.add_paragraph()

    _line(doc, "Acts & Sections: {{ acts_sections_line }}", bold=True)
    doc.add_paragraph()
    _line(doc, "Date & time of seizure: {{ seizure_datetime }}")
    _line(doc, "Place of seizure: {{ seizure_location }}")
    doc.add_paragraph()

    # Person from whom property seized
    _line(doc, "Person from whom property seized:", bold=True)
    _line(doc, "Name: {{ accused_name }}    S/o: {{ accused_father }}    Age: {{ accused_age }}")
    _line(doc, "Address: {{ accused_address }}")
    doc.add_paragraph()

    _line(
        doc,
        "The property described below was seized by {{ io_name }}, Investigating Officer, in "
        "the presence of the two independent panch witnesses named below, and each article was "
        "thereafter packed and sealed:",
    )
    # Seized-property table. (The est.-value column carries its own 'Rs.' prefix, so the
    # header drops the redundant '(Rs.)'.)
    _looped_table(
        doc,
        ["Sr.", "Description of property", "Qty", "Est. value"],
        "{%tr for item in seized_items %}",
        [
            "{{ loop.index }}",
            "{{ item.description }}",
            "{{ item.quantity }}",
            "{{ item.estimated_value }}",
        ],
    )
    doc.add_paragraph()

    # Two independent panch witness blocks
    _line(doc, "Independent panch witnesses:", bold=True)
    _line(doc, "1. Name: {{ witness1.full_name }}    S/o: {{ witness1.father_name }}")
    _line(doc, "   Address: {{ witness1.address }}          Signature: ____________________")
    _line(doc, "2. Name: {{ witness2.full_name }}    S/o: {{ witness2.father_name }}")
    _line(doc, "   Address: {{ witness2.address }}          Signature: ____________________")
    doc.add_paragraph()

    # Seizing / Investigating Officer signature block (rank + buckle number)
    _line(doc, "Seizing / Investigating Officer:", bold=True)
    _line(doc, "Name: {{ io_name }}    Rank: {{ io_rank }}    Buckle (Badge) No.: {{ io_badge_no }}")
    _line(doc, "Police Station: {{ police_station }}          Signature: ____________________")
    doc.add_paragraph()
    _line(doc, "NOTE: Draft generated by CrimeGPT — to be verified and signed by the officer.")
    _apply_fonts(doc)
    doc.save(HERE / "seizure_receipt.docx")


def build_panchnama():
    doc = Document()
    _title(doc, "PANCHNAMA / પંચનામું")
    doc.add_paragraph()
    _line(doc, "Police Station: {{ police_station }}    District: {{ district }}")
    _line(doc, "Case / FIR No.: {{ case_number }}  (FIR No. {{ fir_number }})")
    _line(doc, "Date: {{ panchnama_date }}    Place: {{ panchnama_place }}")
    doc.add_paragraph()
    _line(
        doc,
        "Before me, {{ io_name }}, Investigating Officer of the above case, the panch "
        "witnesses named below were called and the purpose of this panchnama was explained "
        "to them. In their presence the following proceedings were carried out:",
    )
    _line(doc, "Panch witnesses:", bold=True)
    _looped_table(
        doc,
        ["No.", "Name", "Father's name", "Address"],
        "{%tr for w in witnesses %}",
        ["{{ loop.index }}", "{{ w.full_name }}", "{{ w.father_name }}", "{{ w.address }}"],
    )
    doc.add_paragraph()
    _line(doc, "Accused:", bold=True)
    _line(
        doc,
        "{{ accused_name }}, son of {{ accused_father }}, age {{ accused_age }}, "
        "residing at {{ accused_address }}.",
    )
    doc.add_paragraph()
    _line(doc, "Narrative of proceedings:", bold=True)
    _line(doc, "{{ proceedings_narrative }}")
    doc.add_paragraph()
    _line(doc, "Articles found / seized during the proceedings:", bold=True)
    _looped_table(
        doc,
        ["No.", "Description", "Quantity", "Estimated value"],
        "{%tr for item in seized_items %}",
        ["{{ loop.index }}", "{{ item.description }}", "{{ item.quantity }}", "{{ item.estimated_value }}"],
    )
    doc.add_paragraph()
    _line(
        doc,
        "The above panchnama was read over and explained to the panchas, who admitted it "
        "to be correct and signed below.",
    )
    _sig_block(doc, "Panch witnesses (sign)", "{{ io_name }}\nInvestigating Officer")
    _line(doc, "NOTE: Draft generated by CrimeGPT — to be verified and signed by the officer.")
    _apply_fonts(doc)
    doc.save(HERE / "panchnama.docx")


def build_remand_request():
    doc = Document()
    _line(doc, "To,")
    _line(doc, "The Hon'ble Court of the Learned Judicial Magistrate,")
    _line(doc, "{{ district }}.")
    doc.add_paragraph()
    _title(doc, "APPLICATION FOR POLICE CUSTODY REMAND")
    _line(doc, "(Under Section 187 of the Bharatiya Nagarik Suraksha Sanhita, 2023)")
    doc.add_paragraph()
    _line(doc, "Police Station: {{ police_station }}    District: {{ district }}")
    _line(doc, "Case / FIR No.: {{ case_number }}  (FIR No. {{ fir_number }} dated {{ fir_date }})")
    doc.add_paragraph()
    _line(doc, "Respectfully showeth:", bold=True)
    _line(
        doc,
        "1. That the accused {{ accused_name }}, son of {{ accused_father }}, age "
        "{{ accused_age }}, residing at {{ accused_address }}, has been arrested in "
        "connection with the above case.",
    )
    _line(doc, "2. That the following sections of law have been applied in this case:")
    _looped_table(
        doc,
        ["Act", "Section", "Title"],
        "{%tr for s in sections_applied %}",
        ["{{ s.act }}", "{{ s.section_code }}", "{{ s.section_title }}"],
    )
    _line(doc, "3. Investigation carried out so far:")
    _line(doc, "{{ investigation_done }}")
    _line(doc, "4. Investigation still pending:")
    _line(doc, "{{ pending_investigation }}")
    _line(doc, "5. Grounds for seeking police custody:")
    _line(doc, "{{ grounds_for_custody }}")
    doc.add_paragraph()
    _line(
        doc,
        "It is therefore prayed that the accused be remanded to police custody to enable "
        "completion of the investigation.",
    )
    _sig_block(doc, "Place: {{ police_station }}", "{{ io_name }}\nInvestigating Officer")
    _line(doc, "NOTE: Draft generated by CrimeGPT — to be verified and signed by the officer.")
    _apply_fonts(doc)
    doc.save(HERE / "remand_request.docx")


def build_medical_letter():
    doc = Document()
    _line(doc, "To,")
    _line(doc, "The Medical Officer,")
    _line(doc, "{{ hospital }}.")
    doc.add_paragraph()
    _title(doc, "REQUEST FOR MEDICAL EXAMINATION")
    doc.add_paragraph()
    _line(doc, "From: {{ io_name }}, Investigating Officer, {{ police_station }}, {{ district }}.")
    _line(doc, "Case / FIR No.: {{ case_number }}  (FIR No. {{ fir_number }})")
    doc.add_paragraph()
    _line(doc, "Sir / Madam,")
    _line(
        doc,
        "In connection with the above case, {{ subject_name }} ({{ subject_role }}) is being "
        "produced before you. You are requested to conduct the medical examination of the said "
        "person for the following purpose:",
    )
    _line(doc, "Purpose of examination: {{ examination_purpose }}")
    doc.add_paragraph()
    _line(
        doc,
        "You are further requested to furnish the medical examination report and certificate "
        "to this office at the earliest for the purpose of investigation.",
    )
    _sig_block(doc, "Place: {{ police_station }}", "{{ io_name }}\nInvestigating Officer")
    _line(doc, "NOTE: Draft generated by CrimeGPT — to be verified and signed by the officer.")
    _apply_fonts(doc)
    doc.save(HERE / "medical_letter.docx")


# --------------------------------------------------------------------------
# LERS (Law Enforcement Response System) request templates — compliant request
# forms addressed to a platform (Meta / WhatsApp / Instagram). These are TEMPLATES,
# not a live API integration. They use only existing pool context fields; request-
# specific details (target identifier, data sought, date range, exact BNSS section)
# are blank fields the officer fills per request.
def _lers_common_header(doc, title):
    _title(doc, title)
    _line(doc, "COMPLIANT LAW-ENFORCEMENT REQUEST TEMPLATE — not a live platform API integration.", bold=True)
    doc.add_paragraph()
    _line(doc, "To: The Law Enforcement Response Team,")
    _line(doc, "Platform: ____________________  (Meta / WhatsApp / Instagram)")
    doc.add_paragraph()
    _line(doc, "Requesting agency: {{ police_station }}, {{ district }} (India)")
    _line(doc, "Case / FIR No.: {{ case_number }}  (FIR No. {{ fir_number }} dated {{ fir_date }})")
    _line(doc, "Investigating Officer: {{ io_name }}, {{ io_rank }}, Badge/Buckle No. {{ io_badge_no }}")
    doc.add_paragraph()
    _line(doc, "Legal basis: Under Section ________ of the Bharatiya Nagarik Suraksha Sanhita, 2023 (BNSS),")
    _line(doc, "in connection with an offence under {{ acts_sections_line }}.")
    doc.add_paragraph()
    _line(doc, "Target identifier (account ID / phone number / profile URL):", bold=True)
    _line(doc, "________________________________________________________________")
    doc.add_paragraph()


def _lers_footer(doc):
    doc.add_paragraph()
    _sig_block(
        doc,
        "Place: {{ police_station }}\nDate: ____________",
        "{{ io_name }}\n{{ io_rank }} · Badge {{ io_badge_no }}\nInvestigating Officer",
    )
    _line(
        doc,
        "NOTE: Compliant request template generated by CrimeGPT — to be verified and signed "
        "by the officer. This is not a live platform API integration.",
    )


def build_lers_preservation_request():
    doc = Document()
    _lers_common_header(doc, "DATA PRESERVATION REQUEST / માહિતી જાળવણી વિનંતી")
    _line(doc, "Data to be PRESERVED (do not disclose at this stage):", bold=True)
    _line(doc, "  [ ] Subscriber / registration information     [ ] Login IP history")
    _line(doc, "  [ ] Message / call metadata                   [ ] Media and stored content")
    _line(doc, "  Other: ______________________________________________________")
    doc.add_paragraph()
    _line(doc, "Period for which data is to be preserved: from ____________ to ____________")
    doc.add_paragraph()
    _line(
        doc,
        "PRESERVATION STATEMENT: Pursuant to a lawful investigation, you are requested to "
        "PRESERVE and NOT delete or alter the data described above associated with the "
        "identifier(s) above for a period of ninety (90) days, extendable on further request, "
        "pending service of formal legal process for its disclosure. This is a preservation "
        "request only — no disclosure of data is sought at this stage.",
    )
    _lers_footer(doc)
    _apply_fonts(doc)
    doc.save(HERE / "lers_preservation_request.docx")


def build_lers_records_request():
    doc = Document()
    _lers_common_header(doc, "RECORDS DISCLOSURE REQUEST / રેકોર્ડ જાહેર કરવાની વિનંતી")
    _line(doc, "Records / data SOUGHT for disclosure:", bold=True)
    _line(doc, "  [ ] Basic subscriber information              [ ] Login / IP logs")
    _line(doc, "  [ ] Transactional / message metadata         [ ] Stored content (where lawfully permitted)")
    _line(doc, "  Other: ______________________________________________________")
    doc.add_paragraph()
    _line(doc, "Date range of records requested: from ____________ to ____________")
    doc.add_paragraph()
    _line(
        doc,
        "DISCLOSURE STATEMENT: Pursuant to the legal basis stated above, you are requested to "
        "DISCLOSE the records described above for the identifier(s) and period specified, "
        "certified as true records, in a machine-readable format, to the requesting officer "
        "through the platform's Law Enforcement Response System. Please cite the Case / FIR "
        "number above in your response.",
    )
    _lers_footer(doc)
    _apply_fonts(doc)
    doc.save(HERE / "lers_records_request.docx")


def main():
    build_seizure_receipt()
    build_panchnama()
    build_remand_request()
    build_medical_letter()
    build_lers_preservation_request()
    build_lers_records_request()
    print("Built templates in", HERE)
    for f in sorted(HERE.glob("*.docx")):
        print("  -", f.name)


if __name__ == "__main__":
    main()
